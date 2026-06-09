import re
import os
import tempfile
from typing import Dict, Any, List
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Canonical element order used everywhere
_ELEMENT_ORDER = ["subheading", "paragraph", "list_item", "table"]


def _build_output(toc: List[str], title_text: str, title_style: Dict,
                  raw_sections: List[Dict], default_style: Dict) -> Dict[str, Any]:
    """
    Assemble the final lean JSON structure.

    Output shape:
    {
      "table_of_contents": ["Introduction", "Methodology", ...],
      "title": { "text": "...", "style": {...} },
      "sections": [
        {
          "heading_style": {...},
          "element_styles": {          # only types that exist in this section
            "subheading": {...},
            "paragraph":  {...},
            "list_item":  {...},
            "table":      {...}
          },
          "components": ["subheading", "table"]   # per-section, only what exists
        },
        ...
      ],
      "default_style": {...}
    }
    """
    sections_out = []
    for sec in raw_sections:
        # Only include styles for element types actually present
        element_styles = {
            t: sec["styles"][t]
            for t in _ELEMENT_ORDER
            if sec["styles"].get(t)
        }
        components = [t for t in _ELEMENT_ORDER if t in sec["_seen_types"]]

        sections_out.append({
            "heading_style":  sec.get("heading_style", {}),
            "element_styles": element_styles,
            "components":     components,
        })

    return {
        "table_of_contents": toc,
        "title":             {"text": title_text, "style": title_style},
        "sections":          sections_out,
        "default_style":     default_style,
    }


# ── DOCX extractor ────────────────────────────────────────────────────────────

def extract_metadata_from_docx(doc) -> Dict[str, Any]:

    def first_text_run(para: Paragraph):
        for r in para.runs:
            if r.text and r.text.strip():
                return r
        return para.runs[0] if para.runs else None

    def para_style_name(para: Paragraph) -> str:
        try:
            return (para.style.name or "").strip().lower()
        except Exception:
            return ""

    def is_list_paragraph(para: Paragraph, text: str) -> bool:
        sn = para_style_name(para)
        # Catch all common Word list style name variants
        if any(kw in sn for kw in ("list", "bullet", "number", "item")):
            return True
        # Native Word numbering: <w:numPr> marks a list item even when no
        # bullet character appears in para.text (e.g. auto-numbered lists)
        try:
            if para._p.pPr is not None and para._p.pPr.numPr is not None:
                return True
        except Exception:
            pass
        # Visible bullet / number / letter prefix in text
        return bool(re.match(
            r"^(?:[-*•–▪▸►●○◦]"
            r"|\d+[\.)]|[a-zA-Z][\.])\s+",
            text
        ))

    def heading_level(para: Paragraph, text: str, fmt: Dict,
                      body_fs: float, title_fs: float) -> int:
        sn = para_style_name(para)
        m  = re.match(r"heading\s*(\d+)", sn)
        if m:
            lv = int(m.group(1))
            return lv if lv <= 2 else 2
        if text.startswith("##"):
            return 2
        if text.startswith("#"):
            return 1
        if len(text) > 90 or re.search(r"[\.!?]$", text):
            return 0
        size      = float(fmt.get("font_size", 0) or 0)
        bold      = bool(fmt.get("bold"))
        underline = bool(fmt.get("underline"))
        if not size:
            return 0
        large_title = title_fs >= (body_fs + 8.0)
        if (size >= (body_fs + 6.0) and (not large_title or size < title_fs)) \
                or (size >= (body_fs + 2.0) and bold and underline):
            return 1
        if size >= (body_fs + 2.0):
            return 2
        if bold and len(text) < 80 and not text.endswith(('.', '!', '?')):
            return 2
        return 0

    def new_section(heading_style: Dict = None) -> Dict:
        return {
            "heading_style": heading_style or {},
            "styles":        {t: {} for t in _ELEMENT_ORDER},
            "_seen_types":   set(),
        }

    def record(section: Dict, block_type: str, style: Dict = None):
        section["_seen_types"].add(block_type)
        if style and not section["styles"][block_type]:
            section["styles"][block_type] = style

    def table_style(table: Table) -> Dict:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    run = first_text_run(para)
                    if run:
                        return get_run_fmt(run, para, doc)
        return {}

    def iter_blocks(document: Document):
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield "paragraph", Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield "table", Table(child, document)

    # 1. Body font size (weighted by character count)
    fs_counts: Dict[float, int] = {}
    for kind, blk in iter_blocks(doc):
        if kind != "paragraph":
            continue
        for run in blk.runs:
            if run.text and run.text.strip():
                fmt  = get_run_fmt(run, blk, doc)
                size = fmt.get("font_size")
                if size:
                    rs = round(float(size), 1)
                    fs_counts[rs] = fs_counts.get(rs, 0) + len(run.text.strip())
    body_fs = max(fs_counts, key=fs_counts.__getitem__) if fs_counts else 11.0

    # 2. Title (largest font in first 30 non-trivial paragraphs)
    title_text  = ""
    title_style: Dict = {}
    title_fs    = 0.0
    count       = 0
    for kind, blk in iter_blocks(doc):
        if kind != "paragraph":
            continue
        text = blk.text.strip()
        if not text:
            continue
        if len(text) < 30 and (text.isdigit() or "page" in text.lower()
                               or "confidential" in text.lower()):
            continue
        run  = first_text_run(blk)
        fmt  = get_run_fmt(run, blk, doc) if run else {}
        size = float(fmt.get("font_size", 0) or 0)
        if count < 30:
            if size > title_fs:
                title_fs    = size
                title_text  = text
                title_style = fmt
            count += 1
        else:
            break
    if not title_text:
        title_fs    = body_fs + 6.0
        title_text  = "Generated Report"
        title_style = {"font_name": "Calibri", "font_size": title_fs,
                       "bold": True, "italic": False, "underline": False,
                       "color": "000000", "alignment": "left"}

    # 3. Default (body) style
    default_style = {"font_name": "Calibri", "font_size": body_fs, "bold": False,
                     "italic": False, "underline": False, "color": "000000",
                     "alignment": "left"}
    for kind, blk in iter_blocks(doc):
        if kind != "paragraph":
            continue
        text = blk.text.strip()
        if not text or len(text) < 20:
            continue
        run = first_text_run(blk)
        if run:
            fmt  = get_run_fmt(run, blk, doc)
            size = float(fmt.get("font_size", 0) or 0)
            if size <= 9.0:
                continue
            if round(size, 1) == round(body_fs, 1):
                default_style = fmt
                break

    # 4. Walk document → build raw sections
    raw_sections: List[Dict] = []
    current     = None
    toc: List[str] = []

    for kind, blk in iter_blocks(doc):
        if kind == "paragraph":
            para  = blk
            text  = para.text.strip()
            if not text:
                continue
            run   = first_text_run(para)
            if not run:
                continue
            fmt   = get_run_fmt(run, para, doc)
            size  = float(fmt.get("font_size", 0) or 0)
            color = str(fmt.get("color", "")).upper()

            # Drop header/footer/page-number fragments
            if 0 < size <= 9.0:
                tl      = text.lower()
                is_gray = color in {"999999","BBBBBB","CCCCCC","888888","7F7F7F","A0A0A0"}
                is_rt   = fmt.get("alignment") == "right"
                has_pg  = text.isdigit() or "page" in tl \
                          or re.search(r"\b\d+\s*(?:of|/)\s*\d+\b", tl)
                if is_gray or is_rt or has_pg or len(text) < 15:
                    continue

            # Drop document title text
            if title_text == text and not current and not raw_sections:
                continue

            level   = heading_level(para, text, fmt, body_fs, title_fs)
            is_list = is_list_paragraph(para, text)

            if level == 1 and not is_list:
                if current and (current["heading_style"] or current["_seen_types"]):
                    raw_sections.append(current)
                current = new_section(fmt)
                toc.append(text)
                continue

            if current is None:
                current = new_section()

            if level == 2 and not is_list:
                record(current, "subheading", fmt)
            elif is_list:
                record(current, "list_item", fmt)
            else:
                record(current, "paragraph", fmt)

        else:  # table block
            if current is None:
                current = new_section()
            record(current, "table", table_style(blk))

    if current and (current["heading_style"] or current["_seen_types"]):
        raw_sections.append(current)

    return _build_output(toc, title_text, title_style, raw_sections, default_style)


# ── Font-formatting helper ────────────────────────────────────────────────────

def get_run_fmt(run, para, doc) -> Dict[str, Any]:
    fmt = {}

    try:
        normal_style = doc.styles['Normal']
    except Exception:
        normal_style = None

    doc_defaults: Dict[str, Any] = {}
    try:
        rPrDefault = doc.styles.element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')
        if rPrDefault:
            rPr = rPrDefault[0]
            sz  = rPr.xpath('./w:sz/@w:val')
            if sz:
                doc_defaults['size'] = Pt(int(sz[0]) / 2.0)
            rFonts = rPr.xpath('./w:rFonts/@w:ascii')
            if rFonts:
                doc_defaults['name'] = str(rFonts[0])
            if rPr.xpath('./w:b'):
                doc_defaults['bold'] = True
            if rPr.xpath('./w:i'):
                doc_defaults['italic'] = True
            color_el = rPr.xpath('./w:color/@w:val')
            if color_el and color_el[0] != 'auto':
                doc_defaults['color'] = str(color_el[0])
    except Exception:
        pass

    def resolve(prop):
        def valid(v):
            if prop == "color":
                return v is not None and getattr(v, "rgb", None) is not None
            return v is not None
        val = getattr(run.font, prop, None)
        if valid(val):
            return val
        if run.style and run.style.font:
            val = getattr(run.style.font, prop, None)
            if valid(val):
                return val
        style = para.style
        checked_normal = False
        while style:
            if style.name == 'Normal':
                checked_normal = True
            if style.font:
                val = getattr(style.font, prop, None)
                if valid(val):
                    return val
            style = style.base_style
        if not checked_normal and normal_style and normal_style.font:
            val = getattr(normal_style.font, prop, None)
            if valid(val):
                return val
        return doc_defaults.get(prop, None)

    fmt["font_name"] = resolve("name") or "Calibri"
    size = resolve("size")
    fmt["font_size"] = size.pt if size else 11.0
    fmt["bold"]      = resolve("bold")      or False
    fmt["italic"]    = resolve("italic")    or False
    fmt["underline"] = resolve("underline") or False

    color_val = resolve("color")
    if color_val and hasattr(color_val, "rgb") and color_val.rgb:
        fmt["color"] = str(color_val.rgb)
    elif "color" in doc_defaults:
        fmt["color"] = doc_defaults["color"]
    else:
        fmt["color"] = "000000"

    _ALIGN_MAP = {
        WD_ALIGN_PARAGRAPH.CENTER:     "center",
        WD_ALIGN_PARAGRAPH.RIGHT:      "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY:    "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
    }
    align = None
    try:
        align = para.alignment
    except Exception:
        pass
    if align is None:
        try:
            style = para.style
            while style and align is None:
                align = style.paragraph_format.alignment
                style = style.base_style
        except Exception:
            pass
    fmt["alignment"] = _ALIGN_MAP.get(align, "left")
    return fmt


# ── PDF extractor ─────────────────────────────────────────────────────────────

def extract_metadata_from_pdf(pdf_path: str) -> Dict[str, Any]:
    """
    Convert PDF to DOCX via pdf2docx, fix its merged-paragraph output,
    then delegate to extract_metadata_from_docx.

    pdf2docx collapses consecutive same-style lines into one Normal paragraph
    joined by \n.  _split_merged_paragraphs detects bullet/number lines and
    re-injects them as ListParagraph paragraphs so the docx extractor sees them.
    """
    _LIST_RE = re.compile(
        r"^(?:[-*\u2022\u2013\u25aa\u25b8\u25ba\u25cf\u25cb\u25e6]"
        r"|\d+[.)]\s|[a-zA-Z][.]\s)\s*\S"
    )

    def _split_merged_paragraphs(doc):
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        for para in list(doc.paragraphs):
            raw = para.text
            if "\n" not in raw:
                continue
            lines = [ln.strip() for ln in raw.split("\n") if ln.strip()]
            if len(lines) <= 1:
                continue

            parent = para._p.getparent()
            idx    = list(parent).index(para._p)

            for offset, line in enumerate(lines):
                new_p  = copy.deepcopy(para._p)
                # Remove all existing runs
                for r in new_p.findall(f"{{{ns}}}r"):
                    new_p.remove(r)
                # Ensure pPr exists
                pPr = new_p.find(f"{{{ns}}}pPr")
                if pPr is None:
                    pPr = etree.SubElement(new_p, f"{{{ns}}}pPr")
                # Set style
                pStyle = pPr.find(f"{{{ns}}}pStyle")
                if pStyle is None:
                    pStyle = etree.SubElement(pPr, f"{{{ns}}}pStyle")
                style_val = "ListParagraph" if _LIST_RE.match(line) else "Normal"
                pStyle.set(f"{{{ns}}}val", style_val)
                # Add run with line text
                new_r = etree.SubElement(new_p, f"{{{ns}}}r")
                new_t = etree.SubElement(new_r, f"{{{ns}}}t")
                new_t.text = line
                new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                parent.insert(idx + offset, new_p)

            parent.remove(para._p)
        return doc

    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        from pdf2docx import parse
        from lxml import etree
        import copy
        parse(pdf_path, tmp_path, start=0, end=None)
        doc = Document(tmp_path)
        doc = _split_merged_paragraphs(doc)
        result = extract_metadata_from_docx(doc)
    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass
    return result
