import json
import os
import re
import tempfile
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Inches

from models import Block
from config import METADATA_DIR

def detect_structure(plain_text: str) -> List[Block]:
    lines = plain_text.strip().split("\n")
    blocks = []
    pending_paragraph = []

    def is_table_row(line: str) -> bool:
        return "\t" in line or (line.count("|") >= 2 and not line.strip().startswith("-"))

    def flush_paragraph():
        if pending_paragraph:
            blocks.append(Block(type="body", text=" ".join(pending_paragraph)))
            pending_paragraph.clear()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            flush_paragraph()
            i += 1
            continue

        # Table detection
        if is_table_row(line):
            flush_paragraph()
            table_lines = [line]
            i += 1
            while i < len(lines) and is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            for tline in table_lines:
                blocks.append(Block(type="table_row", text=tline))
            continue

        # List detection
        if re.match(r"^(\-|\*|•|–|\d+[\.\)])\s", line):
            flush_paragraph()
            blocks.append(Block(type="list_item", text=line))
            i += 1
            continue

        # Title detectio
        if len(blocks) == 0 and not pending_paragraph:
            blocks.append(Block(type="title", text=line))
            i += 1
            continue

        # Heading detection
        if len(line) <= 70 and not line.endswith("."):
            flush_paragraph()

            # Subheading detection
            if line.startswith("##"):
                blocks.append(Block(type="subheading", text=line.lstrip("#").strip()))
            elif line.startswith("#"):
                blocks.append(Block(type="heading", text=line.lstrip("#").strip()))
            else:
                blocks.append(Block(type="heading", text=line))
            i += 1
            continue

        # Body
        pending_paragraph.append(line)
        i += 1

    flush_paragraph()
    return blocks

def apply_run_fmt(run, fmt):

    if fmt.get("font_name"):
        run.font.name = fmt["font_name"]
    if fmt.get("font_size"):
        run.font.size = Pt(float(fmt["font_size"]))
    if fmt.get("bold") is not None:
        run.bold = bool(fmt["bold"])
    if fmt.get("italic") is not None:
        run.italic = bool(fmt["italic"])
    if fmt.get("underline") is not None:
        val = fmt["underline"]
        if isinstance(val, bool):
            run.underline = val
        elif isinstance(val, int) and val > 0:
            run.underline = True
        else:
            run.underline = False
    if fmt.get("color"):
        try:
            run.font.color.rgb = RGBColor.from_string(str(fmt["color"]))
        except Exception:
            pass

def apply_para_fmt(para, fmt):

    if fmt.get("alignment") is not None:
        para.alignment = int(fmt["alignment"])
    pf = para.paragraph_format
    if fmt.get("line_spacing"):
        pf.line_spacing = float(fmt["line_spacing"])
    if fmt.get("space_before"):
        val = fmt["space_before"]
        pf.space_before = Pt(val / 12700) if isinstance(val, (int, float)) and val > 100 else Pt(val) if val else None
    if fmt.get("space_after"):
        val = fmt["space_after"]
        pf.space_after = Pt(val / 12700) if isinstance(val, (int, float)) and val > 100 else Pt(val) if val else None
    if fmt.get("first_line_indent"):
        val = fmt["first_line_indent"]
        pf.first_line_indent = Pt(val / 12700) if isinstance(val, (int, float)) and val > 100 else Pt(val) if val else None

def generate_docx(template_id: str, blocks: List[Block]) -> str:

    # Loading metadata
    meta_path = os.path.join(METADATA_DIR, f"{template_id}.json")
    if not os.path.exists(meta_path):
        raise FileNotFoundError(f"Template '{template_id}' not found")
    with open(meta_path) as f:
        metadata = json.load(f)
    sections_meta = metadata.get("sections", [])
    basic_style = metadata.get("basic_style", {})
    title_style = metadata.get("title", {}).get("style", {})

    first_section = sections_meta[0] if sections_meta else {}
    section_styles = first_section.get("styles", {}) if isinstance(first_section, dict) else {}
    heading_style = first_section.get("heading_style", {}) if isinstance(first_section, dict) else {}

    fmt_map = {
        "title": title_style or heading_style or basic_style,
        "heading": heading_style or title_style or basic_style,
        "subheading": section_styles.get("subheading", {}) or heading_style or basic_style,
        "body": section_styles.get("paragraph", {}) or basic_style,
        "paragraph": section_styles.get("paragraph", {}) or basic_style,
        "list_item": section_styles.get("list_item", {}) or section_styles.get("paragraph", {}) or basic_style,
        "table_row": section_styles.get("table", {}) or section_styles.get("paragraph", {}) or basic_style,
    }

    doc = Document()

    # Applying basic page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)

    for block in blocks:
        fmt = fmt_map.get(block.type, fmt_map.get("body", basic_style))

        if block.type == "table_row":
            cells = [c.strip() for c in block.text.split("\t") if c.strip()]
            if not cells:
                cells = [c.strip() for c in block.text.split("|") if c.strip()]
            if not cells:
                continue
            table = doc.add_table(rows=1, cols=len(cells), style="Table Grid")
            for idx, cell_text in enumerate(cells):
                cell = table.cell(0, idx)
                cell.text = cell_text
                for run in cell.paragraphs[0].runs:
                    apply_run_fmt(run, fmt)
            doc.add_paragraph()
        elif block.type == "list_item":
            para = doc.add_paragraph()

            text = re.sub(r"^(\-|\*|•|–|\d+[\.\)])\s*", "", block.text)
            run = para.add_run(f"• {text}")
            apply_run_fmt(run, fmt)
            apply_para_fmt(para, fmt)
        else:
            para = doc.add_paragraph()
            run = para.add_run(block.text)
            apply_run_fmt(run, fmt)
            apply_para_fmt(para, fmt)

    # Saving to temp file and returning
    out_path = tempfile.mktemp(suffix=".docx")
    doc.save(out_path)
    return out_path
